# !/usr/bin/env python
# -*- coding:utf8 -*- 

# **************************************************************************
# * Copyright (c) 2024. All Rights Reserved
# **************************************************************************
# * @function OCR Demo, tesseract+Vite UI 实现, windows 环境
# * @author wqw547243068@163.com
# * @date 2024/12/05 17:00
# **************************************************************************


import os
import json
import logging
import threading
from queue import Queue
import socket
import time
import random
from flask import Flask, request, send_file
from conf import response_info, font_path

app = Flask(__name__)

# 日志系统配置
handler = logging.FileHandler('app.log', encoding='UTF-8') # 日志输出到终端
# handler = logging.StreamHandler() # 日志输出到控制台
#设置日志文件，和字符编码
logging_format = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(filename)s - %(funcName)s - %(lineno)s - %(message)s')
handler.setFormatter(logging_format)
app.logger.addHandler(handler)
#设置日志存储格式，也可自定义日志格式满足不同的业务需求
logger = logging.getLogger(__name__)

from PIL import Image
from docx import Document
# import pdfplumber
# import pytesseract
from paddleocr import PaddleOCR, draw_ocr
# import langid


cur_dir = os.path.dirname(__file__)
cache_dir = os.path.join(cur_dir, os.pardir, 'cache_file')
# ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = cache_dir
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 设置最大文件上传大小为 100MB

if not os.path.exists(cache_dir):
    os.mkdir(cache_dir)

# 多语种支持
languages = ['ch', 'en', 'japan', 'fr', 'de']  # 需要支持的语言列表

# 模型初始化
# ocr = PaddleOCR(use_angle_cls=True, lang='ch')
ocr_api_info = {}
for lang in languages:
    ocr_api_info[lang] = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)


def extract_ip():
    st = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:       
        st.connect(('10.255.255.255', 1))
        IP = st.getsockname()[0]
    except Exception:
        IP = '127.0.0.1'
    finally:
        st.close()
    return IP

# 服务器 ip地址: 跨机器访问时图片无法显示，ip地址从localhost改127.0.0.1或0.0.0.0都不管用，最后用实际ip地址
# http://0.0.0.0:5000/download?fileId=123.tar
local_host_ip = extract_ip()
local_file_url = f'http://{local_host_ip}:5000/download?fileId='

def log(msg):
    """
        多进程日志输出
    """
    pid = os.getpid()
    tid = threading.current_thread().ident
    logging.info(f"进程[{pid}]-线程[{tid}]: {msg}")

def getResult(lock, lang, file_name, q):
    """
        单次 OCR 请求
    """
    log('开始请求OCR服务')
    # 开始请求OCR服务
    result = ocr_api_info[lang].ocr(file_name)
    # 多种返回结果: 空、一个识别结果、多个识别结果
    # [null]
    #  [[
    #     [[[1931.0, 1453.0], [1990.0, 1453.0], [1990.0, 1487.0], [1931.0, 1487.0]], ["191t", 0.7493318915367126]]
    #  ]]
    # [[
    #     [[[297.0, 157.0], [1755.0, 124.0], [1760.0, 337.0], [301.0, 370.0]], ["很多人不需要再见!", 0.9498765468597412]], 
    #     [[[351.0, 493.0], [1773.0, 468.0], [1776.0, 657.0], [354.0, 683.0]], ["因为只是路过而已.", 0.8729634881019592]], 
    #     [[[334.0, 833.0], [1755.0, 842.0], [1753.0, 1049.0], [333.0, 1039.0]], ["遗忘就是我给你", 0.9973150491714478]], 
    #     [[[404.0, 1153.0], [1457.0, 1182.0], [1450.0, 1434.0], [397.0, 1404.0]], ["最好的纪念。", 0.951411783695221]], 
    #     [[[1935.0, 1460.0], [1988.0, 1460.0], [1988.0, 1485.0], [1935.0, 1485.0]], ["19楼", 0.9435752034187317]]
    # ]]
    lock.acquire()
    if not result[0]:
        score_avg = 0
    else:
        # 计算平均得分
        score_list = [i[1][1] for i in result[0]]
        score_avg = sum(score_list)/len(score_list)
        print(f'[Note] {lang=}: \t{score_avg}\t{json.dumps([i[1][0] for i in result[0]], ensure_ascii=False)}')
    lock.release()
    q.put([lang, score_avg, result[0]])
    log('请求完毕')



def parseDoc(file_name):
    """
        word 文档解析
    """
    content_info = {"num": 0, "content":[], "status":0, 'msg':'-', 'merge_image':'-'}
    # 读取文档
    try:
        doc = Document(file_name)
        print(f"文件{file_name}解析成功")
        content_info['num'] = len(doc.paragraphs)
        for paragraph in doc.paragraphs:
            content_info['content'].append(paragraph.text)
        content_info['status'] = 1
        content_info['msg'] = 'word解析完毕'
    except Exception as err:
        # logging.error(f"word文档解析失败 {file_name} ...")
        content_info['status'] = -2
        content_info['msg'] = err
    return content_info

def parsePDF(file_name):
    """
        pdf 文档解析
    """
    content_info = {"num": 0, "content":[], "status":0, "msg":'-', 'merge_image':'-'}

    # 读取文档
    try:
        # pdf =  pdfplumber.open(file_name)
        # content_info['num'] = len(pdf.pages)
        # # for page in pdf.pages:
        # #     content_info['content'].append(page.extract_text())
        # for i, page in enumerate(pdf.pages):
        #     content_info['content'].extend([f'==[第{i}页]==', page.extract_text()])

        result = ocr_api_info['ch'].ocr(file_name, cls=True)
        content_list = ['识别结果:']
        for idx, res in enumerate(result):
            content_list.extend([f'\n【第{idx+1}页】']+[i[1][0] for i in res]+['-'*30])
        content_info['content'] = content_list
        content_info['status'] = 1
        content_info['msg'] = 'pdf解析完毕'
    except Exception as err:
        # logging.error(f"pdf文档解析失败 {file_name} ...")
        content_info['status'] = -1
        content_info['msg'] = err
    return content_info


def parseOCR(file_name, hand=False):
    """
        图片OCR: tesseract 
        特点: 支持多语种, 但手写体、场景图片文字识别效果一般
    """
    content_info = {"num": 0, "content":[], "status":0, "msg":'-', 'merge_image':'-'}
    # 读取文档
    try:
        if hand:
            out = []
            # out = pytesseract.image_to_string(file_name, lang='chi_sim', config="--psm 4") # 中文手写体识别
        else:
            out = []
            # out = pytesseract.image_to_string(file_name, lang='chi_sim+eng+deu+fra+rus+jpn')
        # 数据示例： g 611 283 631 297 0
        out = out.split('\n') if out else []
        # out = [ i.replace(' ','') for i in out.split('\n')]
        content_info['num'] = len(out)
        content_info['content'] = out
        content_info['status'] = 1
        content_info['msg'] = '图片OCR解析完毕'
    except Exception as err:
        # logging.error(f"pdf文档解析失败 {file_name} ...")
        content_info['status'] = -1
        content_info['msg'] = err
    return content_info


def parseOCRNew(file_name):
    """
        图片OCR: PaddleOCR,
        特点: 识别效果更好，但不支持多语种自动检测
    """
    content_info = {"num": 0, "content":[], "status":0, "msg":'-', "merge_image":'-'}
    # 读取文档
    try:
        out = ocr_api_info['ch'].ocr(file_name, cls=True) # 方向分类器
        # result = ocr.ocr(file_name, det=False) # 每个item只有文本内容和置信度
        # result = ocr.ocr(file_name, cls=True, det=False) # 不需要文本框，每个item只有文本内容和置信度
        # result = ocr.ocr(file_name, cls=True, rec=False) # 不需要文本内容，每个item只有文本框
        # result = ocr.ocr(file_name, cls=True, rec=False, det=False) #  仅执行方向分类器, 返回分类结果+置信度
        # 数据格式样例, [边框点坐标, [识别文本, 置信度]]
        # [[[442.0, 173.0], [1169.0, 173.0], [1169.0, 225.0], [442.0, 225.0]], ['ACKNOWLEDGEMENTS', 0.99283075]]

        # 结果可视化展示
        result = out[0]
        image = Image.open(file_name).convert('RGB')
        boxes, txts, scores = [], [], []
        for line in result:
            if not line:
                continue
            boxes.append(line[0])
            txts.append(line[1][0])
            scores.append(line[1][1])

        im_show = draw_ocr(image, boxes, txts, scores, font_path=font_path)
        im_show = Image.fromarray(im_show)
        new_file = os.path.basename(file_name).replace('.', '_merge.')
        remote_file = local_file_url + new_file
        im_show.save(os.path.join(os.path.dirname(file_name), new_file))
        # im_show.show('result.jpg')

        # out = out.split('\n') if out else []
        # out = [ i.replace(' ','') for i in out.split('\n')]
        content_info['num'] = len(out)
        content_info['content'] = [i[1][0] for i in result]
        content_info['status'] = 1
        content_info['msg'] = '图片OCR二次解析完毕'
        content_info['merge_image'] = remote_file
    except Exception as err:
        # logging.error(f"pdf文档解析失败 {file_name} ...")
        content_info['status'] = -1
        content_info['msg'] = err
    return content_info


def multiOCR(file_name):
    """
        图片OCR: PaddleOCR, 多进程识别，选最优
        特点: 识别效果更好，但不支持多语种自动检测
    """
    content_info = {"num": 0, "content":[], "status":0, "msg":'-', "merge_image":'-'}

    thread_lock = threading.Lock()

    job_list = []
    q = Queue() # 存储结果

    for lang in languages:
        job = threading.Thread(target=getResult, args=(thread_lock, lang, file_name, q), name=f'job_{lang}')
        job.start()
        job_list.append(job)
    
    # 阻塞在主进程前面
    for thread in job_list:
        thread.join()

    results = []
    for _ in languages:
        # [lang, score_avg, result[0]]
        results.append(q.get())
    
    # 按得分降序排列
    sort_result = sorted(results, key=lambda x: x[1], reverse=True)
    merge_list = []
    for res in sort_result:
        if res[1] < 0.7:
            continue
        merge_list.extend([f'[语种]{res[0]}', f'[得分]{res[1]}']+[i[1][0] for i in res[2]])

    # best_result = max(results, key=lambda x: x[1])
    best_result = sort_result[0]
    if not best_result: # 结果为空
        best_result = ['未知', 0, '检测结果为空']
    text = '\n'.join([i[1][0] for i in best_result[2]])
    print(f'Best Result: {best_result[0]}\t{best_result[1]}\t{best_result[2]}')
    print('Result: ', json.dumps(text, ensure_ascii=False))

    # 结果可视化展示
    result = best_result[2]
    image = Image.open(file_name).convert('RGB')
    boxes, txts, scores = [], [], []
    for line in result:
        if not line:
            continue
        boxes.append(line[0])
        txts.append(line[1][0])
        scores.append(line[1][1])

    im_show = draw_ocr(image, boxes, txts, scores, font_path=font_path)
    im_show = Image.fromarray(im_show)
    new_file = os.path.basename(file_name).replace('.', '_merge.')
    remote_file = local_file_url + new_file
    im_show.save(os.path.join(os.path.dirname(file_name), new_file))
    # im_show.show('result.jpg')

    content_info['num'] = len(result)
    content_info['content'] = merge_list
    # content_info['content'] = [f'[语种]{best_result[0]}', f'[得分]{best_result[1]}']+[i[1][0] for i in result]
    content_info['status'] = 2
    content_info['msg'] = '图片OCR并行解析完毕'
    content_info['merge_image'] = remote_file

    return content_info

@app.route('/')
def home():
    app.logger.info('主页请求')
    # app.logger.info('message info is %s', message, exc_info=1)。
    #app.logger.exception('%s', e) # 异常信息
    return 'OCR接口主页!'
    # return render_template('index.html') 


@app.route("/download")
def download_file():
    """
    下载 src_file 目录下的文件
    eg: 下载当前目录下123.tar 文件 http://0.0.0.0:5000/download?fileId=123.tar
    """
    file_name = request.args.get('fileId')
    file_path = os.path.join(cache_dir, file_name)
    if os.path.isfile(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "The downloaded file does not exist"

@app.route('/api_ocr', methods = ["GET", "POST"])
def post_data():
    # 获取客户端的文件信息
    remote_file = '-'
    # ======= 统一 ======
    if request.method == 'POST':
        # data = request.json
        # data = request.form.to_dict()
        # data = request.values
        data = request.files
        file = data['uploadFile'] # 请求方设置的文件名字段
        # 文件写入磁盘
        cur_file = os.path.join(cache_dir, file.filename)
        file.save(cur_file)
        remote_file = local_file_url + file.filename
    elif request.method == 'GET':
        data = request.args
        cur_file = data['uploadFile']
        remote_file = cur_file
    
    logger.info(f"请求参数: {data=}, {cur_file=}")
    res = {"status":0, 
            "msg":'-', 
            "data":{
                "content":[], # 识别内容
                "merge_image": remote_file, # 融合图
                "scores":[], # 分值
            }, 
            'req':{'file_name': cur_file} # 请求信息
    }
    
    if not os.path.exists(cur_file):
        res['status'] = -1
        res['msg'] = f"文件{cur_file}不存在"
        logger.error(res['msg'])
        return json.dumps(res, ensure_ascii=False)

    res['req']['file_name'] = os.path.basename(cur_file)
    if cur_file.find('.') != -1:
        cur_ext = cur_file.split('.')[-1]
    else:
        cur_ext = '-'
    cur_ext = cur_ext.lower()
    res['data']['file_type'] = cur_ext
    
    # 预设字典
    file_ext = res['req']['file_name']
    if file_ext in response_info:
        res['data']['content'] = response_info[file_ext]['content'].split('\n')
        res['data']['num'] = len(res['data']['content'])
        res['status'] = 1
        res['msg'] = '字典映射'
        res['data']['merge_image'] = local_file_url + response_info[file_ext]['merge_image']
        sleeptime = random.uniform(1, 2)
        time.sleep(sleeptime)
        return res
    
    # 格式检测
    if cur_ext in ('doc', 'docx'): # word 文档, 直接解析
        res['data']['content'] = ['word 文档内容']
        res['msg'] = 'word文件'
        out = parseDoc(cur_file)
        res['status'] = out['status']
        if out['merge_image'] == '-':
            res['data']['merge_iamge'] = local_file_url + 'word.jpg'
        if out['msg'] != '-':
            res['msg'] = out['msg']
        if out['status'] > 0:
            res['data']['content'] = out['content']
        else:
            logger.error(f"word文档解析失败 {cur_file=} -> {out=} ...")
    elif cur_ext in ('pdf'): # pdf 文档
        res['data']['content'] = ['pdf 文档内容']
        res['msg'] = 'pdf文件'
        out = parsePDF(cur_file)
        res['status'] = out['status']
        if out['merge_image'] == '-':
            res['data']['merge_image'] = local_file_url + 'pdf.jpg'
        else:
            res['data']['merge_image'] = local_file_url + out['merge_image']
        if out['msg'] != '-':
            res['msg'] = out['msg']
        if out['status'] > 0:
            res['data']['content'] = out['content']
        else:
            logger.error(f"pdf文档解析失败 {cur_file=} -> {out=}...")
    elif cur_ext in ('jpg', 'png', 'jpeg', 'tiff', 'bmp', 'gif'):
        # 调用ocr工具
        # res = some_function(cur_file)
        res['data']['content'] = ['ocr 文档内容']
        res['msg'] = '图片文件'
        # out = parseOCR(cur_file)
        out = multiOCR(cur_file)
        # # 判断识别出来的文本质量, 语种检测非中文时, 重新生成
        # detect_res = False # 判断是否中文
        # if out['num'] > 0:
        #     tmp = langid.classify(','.join(out['content']))
        #     if tmp[0] == 'zh':
        #         detect_res = True
        # if out['num'] == 0 or not detect_res:
        #     # print(f'二次检测, {out}')
        #     logger.warning('启动二次检测, 疑似手写体')
        #     # 启动中文手写体识别
        #     # out = parseOCR(cur_file, hand=True)
        #     out = parseOCRNew(cur_file)
        res['status'] = out['status']
        if out['msg'] != '-':
            res['msg'] = out['msg']
        if out['status'] > 0:
            res['data']['content'] = out['content'] if out['num'] > 0 else '[未识别到内容]'
            # 生成合成图
            if out['merge_image'] != '-':
                res['data']['merge_image'] = out['merge_image']
        else:
            res['data']['content'] = f'图片解析失败 {cur_file=} -> {out=}...'
            logger.error(f"图片解析失败 {cur_file=} -> {out=}...")
    else:
        res['data']['content'] = ['非预期格式']
        res['msg'] = '其它文件'
        out = []
    print(f"请求参数: {data=}, {out=}, {res=}")
    # return res
    # return jsonify(res)
    return json.dumps(res, ensure_ascii=False)
    # return ';'.join(res['data']['content'])


if __name__ == '__main__':
    #app.run() # 本地访问：只能从自己的计算机上访问
    app.run(host='0.0.0.0', debug=True) # 外网可访问

