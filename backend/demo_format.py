# !/usr/bin/env python
# -*- coding:utf8 -*- 

# **************************************************************************
# * Copyright (c) 2024. All Rights Reserved
# **************************************************************************
# * @function OCR Demo, word/pdf文件解析, windows 环境
# * @author wqw547243068@163.com
# * @date 2024/12/05 17:00
# **************************************************************************


import os
from docx import Document

main_dir = "e:\ocr\data\联合国宣言"
read_path = f"{main_dir}\中英.docx"

# 读取文档
doc = Document(read_path)

print(f"[doc] ----- 读取 ------")
# （1）段落
print('段落句柄: ', len(doc.paragraphs)) # 段落输出
# 输出列表，一共有4份内容
# [<docx.text.paragraph.Paragraph object at 0x7fca95f0aba8>,
# <docx.text.paragraph.Paragraph object at 0x7fca95f0abe0>,
# <docx.text.paragraph.Paragraph object at 0x7fca95f0ab70>, 
#<docx.text.paragraph.Paragraph object at 0x7fca95f0ac50>,]

for paragraph in doc.paragraphs:
    print(paragraph.text)

print(f"[doc] ----- 写入 ------")
save_file = f"{main_dir}\..\\t.docx"
if not os.path.exists(save_file):
    print(f'文件不存在, 创建... {save_file}')
    f = open(save_file, 'w')
    # f.write('新文件')
    f.close()

doc_save = Document()

doc_save.add_heading('文档标题', 0)
doc_save.add_heading('一级标题', 1)
paragraph3 = doc_save.add_paragraph() # 创建空段落，用于填充文字块
# 文字块格式化格式化
paragraph3.add_run("这句是被加粗了文字块").bold = True # 粗体
paragraph3.add_run("，这句是普通文字块，")
paragraph3.add_run("这句是斜体文字块").italic = True # 斜体

doc_save.save(save_file)

print('保存完毕')

print(f"[pdf] ----- 读取 ------")
# pip install pdfplumber

import pdfplumber

read_path = f"{main_dir}\中英.pdf"
# 读取pdf文件，保存为pdf实例
pdf =  pdfplumber.open(read_path) 
# 通过pdfplumber.PDF类的metadata属性获取pdf信息
print(pdf.metadata)
# 通过pdfplumber.PDF类的metadata属性获取pdf页数
print(len(pdf.pages))
first_page = pdf.pages[0] # 第一页pdfplumber.Page实例
print('页码：', first_page.page_number) # 查看页码
print('页宽：', first_page.width) # 查看页宽
print('页高：', first_page.height) # 查看页高
# 读取文本
text = first_page.extract_text()
print(text)